#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
This is the GoReport class. GoReport handles everything from connecting to the target Gophish
server to pulling campaign information and reporting the results.
"""

try:
    # 3rd Party Libraries
    from gophish import Gophish
except:
    print("[!] Could not import the Gophish library! Make sure it is installed.\n\
Run: `python3 -m pip intall gophish`\n\
Test it by running `python3` and then, in the \
Python prompt, typing `from gophish import Gophish`.")
    exit()

# Standard Libraries
import configparser
import os.path
import sys
from collections import Counter
import collections
from datetime import datetime
import re
import pandas as pd


# 3rd Party Libraries
import requests
import xlsxwriter
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor, Inches
from requests.packages.urllib3.exceptions import InsecureRequestWarning
from user_agents import parse
import docx
import random
import matplotlib.pyplot as plt
import matplotlib.dates as md
import nmap
from os import listdir
from os.path import isfile, join
import seaborn as sns

requests.packages.urllib3.disable_warnings(InsecureRequestWarning)


class Goreport(object):
    """
    This class uses the Gophish library to create a new Gophish API connection
    and queries Gophish for information and results related to the specified
    campaign ID(s).
    """
    # Name of the config file -- default is ``gophish.config``
    goreport_config_file = "gophish.config"
    verbose = False

    # Variables for holding Gophish models
    results = None
    campaign = None
    timeline = None

    # Variables for holding campaign information
    cam_id = None
    cam_url = None
    cam_name = None
    cam_status = None
    launch_date = None
    created_date = None
    cam_page_name = None
    cam_smtp_host = None
    completed_date = None
    cam_redirect_url = None
    cam_from_address = None
    cam_subject_line = None
    cam_template_name = None
    cam_capturing_passwords = None
    cam_capturing_credentials = None

    # Variables and lists for tracking event numbers
    total_sent = 0
    total_opened = 0
    total_targets = 0
    total_clicked = 0
    total_reported = 0
    total_submitted = 0
    total_unique_opened = 0
    total_unique_clicked = 0
    total_unique_reported = 0
    total_unique_submitted = 0
    targets_opened = []
    targets_clicked = []
    targets_reported = []
    targets_submitted = []

    # Lists and dicts for holding prepared report data
    campaign_results_summary = []

    # Lists for holding totals for statistics
    browsers = []
    locations = []
    ip_addresses = []
    ip_and_location = {}
    operating_systems = []

    browsers_family = []
    operating_systems_family =[]

    time_df = pd.DataFrame(columns=("time","count","event"))

    # Output options
    report_format = None
    output_word_report = None
    output_xlsx_report = None
    xlsx_header_bg_color = "#0085CA"
    xlsx_header_font_color = "#FFFFFF"

    def __init__(self, report_format, config_file, google, verbose, company):
        """
        Initiate the connection to the Gophish server with the provided host, port,
        and API key and prepare to use the external APIs.
        """
        try:
            # Check if an alternate config file was provided
            if config_file:
                self.goreport_config_file = config_file
            # Open the config file to make sure it exists and is readable
            config = configparser.ConfigParser()
            config.read(self.goreport_config_file)                                                                  # liest Config-File als ein richtiges Config-format ein mit configparser
        except Exception as e:
            print(f"[!] Could not open {self.goreport_config_file} -- make sure it exists and is readable.")
            print(f"L.. Details: {e}")
            sys.exit()                                                                                              # beendet den Prozess, weil config-datei nicht existiert oder nicht gelesen werden kann

        try:
            # Read in the values from the config file
            GP_HOST = self.config_section_map(config, 'Gophish')['gp_host']
            API_KEY = self.config_section_map(config, 'Gophish')['api_key']
        except Exception as e:
            print("[!] There was a problem reading values from the gophish.config file!")
            print(f"L.. Details: {e}")
            sys.exit()                                                                                              # beendet den Prozess, weil inhalt in config-datei nicht eingelsen werden kann

        try:
            # Read in the values from the config file
            self.IPINFO_TOKEN = self.config_section_map(config, 'ipinfo.io')['ipinfo_token']                        # einlesen der config infos aus config datei für ipinfo.io -> nicht unbedingt notwendig
            if not self.IPINFO_TOKEN:
                self.IPINFO_TOKEN = None
        except Exception as e:
            self.IPINFO_TOKEN = None
            print("[!] No ipinfo.io API token was found in the config. GoReport will not lookup IP addresses with ipinfo.io for additional location data.")
            print(f"L.. Details: {e}")

        try:
            # Read in the values from the config file
            self.GEOLOCATE_TOKEN = self.config_section_map(config, 'Google')['geolocate_key']                       # einlesen der config infos aus config datei für geolocate -> nicht unbedingt notwendig
            if not self.GEOLOCATE_TOKEN:
                self.GEOLOCATE_TOKEN = None
        except Exception as e:
            self.GEOLOCATE_TOKEN = None
            if google:
                print("[!] No Google Maps API token was found in the config so GoReport will ignore the `--google` flag.")
                print(f"L.. Details: {e}")

        # Set command line options for the GoReport object
        self.google = google
        self.verbose = verbose
        self.report_format = report_format
        self.company = company
        # Connect to the Gophish API
        # NOTE: This step succeeds even with a bad API key, so the true test is fetching an ID
        print(f"[+] Connecting to Gophish at {GP_HOST}")
        print(f"L.. The API Authorization endpoint is: {GP_HOST}/api/campaigns/?api_key={API_KEY}")
        self.api = Gophish(API_KEY, host=GP_HOST, verify=False)                                                     # Verbindung zu GoPhish API wird hergestellt

    def run(self, id_list, combine_reports, set_complete_status):
        """Run everything to process the target campaign."""
        # Output some feedback for user options
        if combine_reports:
            print("[+] Campaign results will be combined into a single report.")
        if set_complete_status:
            print('[+] Campaign statuses will be set to "Complete" after processing the results.')
        try:
            # Create the list of campaign IDs
            temp_id = []
            # Handle a mixed set of ranges and comma-separated IDs
            if "-" and "," in id_list:
                temp = id_list.split(",")
                for x in temp:
                    if "-" in x:
                        lower = x.split("-")[0]
                        upper = x.split("-")[1]
                        for y in range(int(lower), int(upper) + 1):
                            temp_id.append(str(y))
                    else:
                        temp_id.append(x)
            # Process IDs provided as one or more ranges
            elif "-" in id_list:
                lower = id_list.split("-")[0]
                upper = id_list.split("-")[1]
                for y in range(int(lower), int(upper) + 1):
                    temp_id.append(str(y))
            # Process single or only comma-separated IDs
            else:
                temp_id = id_list.split(",")
            id_list = temp_id
        except Exception as e:
            print("[!] Could not interpret your provided campaign IDs. \
Ensure the IDs are provided as comma-separated integers or interger ranges, e.g. 5,50-55,71.")
            print(f"L.. Details: {e}")
            sys.exit()
        # Begin processing the campaign IDs by removing any duplicates
        try:
            # Get length of user-provided list
            initial_len = len(id_list)
            # Remove duplicate IDs and sort IDs as integers
            id_list = sorted(set(id_list), key=int)
            # Get length of unique, sorted list
            unique_len = len(id_list)
        except Exception as e:
            temp = []
            for id in id_list:
                try:
                    int(id)
                except:
                    temp.append(id)
            print(f"[!] There are {len(temp)} invalid campaign ID(s), i.e. not an integer.")
            print(f"L.. Offending IDs: {','.join(temp)}")
            print(f"L.. Details: {e}")
            sys.exit()
        print(f"[+] A total of {initial_len} campaign IDs have been provided for processing.")
        # If the lengths are different, then GoReport removed one or more dupes
        if initial_len != unique_len:
            dupes = initial_len - unique_len
            print(f"L.. GoReport found {dupes} duplicate campaign IDs, so those have been trimmed.")
        # Provide  list of all IDs that will be processed
        print(f"[+] GoReport will process the following campaign IDs: {','.join(id_list)}")
        # If --combine is used with just one ID it can break reporting, so we catch that here
        if len(id_list) == 1 and combine_reports:
            combine_reports = False
        # Go through each campaign ID and get the results
        campaign_counter = 1
        for CAM_ID in id_list:
            print(f"[+] Now fetching results for Campaign ID {CAM_ID} ({campaign_counter}/{len(id_list)}).")
            try:
                # Request the details for the provided campaign ID
                self.campaign = self.api.campaigns.get(campaign_id=CAM_ID)
            except Exception as e:
                print(f"[!] There was a problem fetching this campaign {CAM_ID}'s details. Make sure your URL and API key are correct. Check HTTP vs HTTPS!")
                print(f"L.. Details: {e}")
            try:
                try:
                    # Check to see if a success message was returned with a message
                    # Possible reasons: campaign ID doesn't exist or problem with host/API key
                    if self.campaign.success is False:
                        print(f"[!] Failed to get results for campaign ID {CAM_ID}")
                        print(f"L.. Details: {self.campaign.message}")
                        # We can't let an error with an ID stop reporting, so check if this was the last ID
                        if CAM_ID == id_list[-1] and combine_reports:
                            self.generate_report()
                # If self.campaign.success does not exist then we were successful
                except:
                    print("[+] Success!")
                    # Collect campaign details and process data
                    self.collect_all_campaign_info(combine_reports)
                    self.process_timeline_events(combine_reports)
                    self.process_results(combine_reports)
                    # If the --complete flag was set, now set campaign status to Complete
                    if set_complete_status:
                        print(f"[+] Setting campaign ID {CAM_ID}'s status to Complete.")
                        try:
                            set_complete = self.api.campaigns.complete(CAM_ID)
                            try:
                                if set_complete.success is False:
                                    print(f"[!] Failed to set campaign status for ID {CAM_ID}.")
                                    print(f"L.. Details: {set_complete.message}")
                            # If set_complete.success does not exist then we were successful
                            except:
                                pass
                        except Exception as e:
                            print(f"[!] Failed to set campaign status for ID {CAM_ID}.")
                            print(f"L.. Details: {e}")
                    # Check if this is the last campaign ID in the list
                    # If this is the last ID and combined reports is on, generate the report
                    if CAM_ID == id_list[-1] and combine_reports:
                        self.generate_report()
                    # Otherwise, if we are not combining reports, generate the reports
                    elif combine_reports is False:
                        self.generate_report()
                    campaign_counter += 1
            except Exception as e:
                print(f"[!] There was a problem processing campaign ID {CAM_ID}!")
                print(f"L.. Details: {e}")
                sys.exit()

    def lookup_ip(self, ip):
        """Lookup the provided IP address with ipinfo.io for location data.

        Example Result:
            {'ip': '52.44.93.197',
            'hostname': 'ec2-52-44-93-197.compute-1.amazonaws.com',
            'city': 'Beaumont',
            'region': 'Texas',
            'country': 'US',
            'loc': '30.0866,-94.1274',
            'postal': '77702',
            'phone': '409',
            'org': 'AS14618 Amazon.com, Inc.'}
        """
        ipinfo_url = f"https://ipinfo.io/{ip}?token={self.IPINFO_TOKEN}"
        try:
            r = requests.get(ipinfo_url)
            return r.json()
        except Exception as e:
            print(f"[!] Failed to lookup `{ip}` with ipinfo.io.")
            print(f"L.. Details: {e}")
            return None

    def get_google_location_data(self, lat, lon):
        """Use Google's Maps API to collect location info for the provided latitude and longitude.

        Google returns a bunch of JSON with a variety of location data. This function returns
        Google's pre-formatted `formatted_address` key for a human-readable address.
        """
        google_maps_url = f"https://maps.googleapis.com/maps/api/geocode/json?latlng={lat},{lon}&sensor=false&key={self.GEOLOCATE_TOKEN}"
        r = requests.get(google_maps_url)
        maps_json = r.json()
        if r.ok:
            try:
                if "error_message" in maps_json:
                    print(f"[!] Google Maps returned an error so using Gophish coordinates. Error: {maps_json['error_message']}")
                    return f"{lat}, {lon}"
                first_result = maps_json['results'][0]
                if "formatted_address" in first_result:
                    return first_result["formatted_address"]
                # In case that key is ever unavailable try to assemble an address
                else:
                    components = first_result['address_components']
                    country = town = None
                    for c in components:
                        if "country" in c['types']:
                            country = c['long_name']
                        if "locality" in c['types']:
                            town = c['long_name']
                        if "administrative_area_level_1" in c['types']:
                            state = c['long_name']
                    return f"{town}, {state}, {country}"
            except Exception as e:
                print("[!] Failed to parse Google Maps API results so using Gophish coordinates.")
                print(f"L.. Error: {e}")
                return f"{lat}, {lon}"
        else:
            print(f"[!] Failed to contact the Google Maps API so using Gophish coordinates. Status code: {r.status_code}")
            return f"{lat}, {lon}"

    def geolocate(self, target, ipaddr, google=False):
        """Attempt to get location data for the provided target and event. Will use ipinfo.io if an
        API key is configured. Otherwise the Gophish latitude and longitude coordinates will be
        returned. If `google` is set to True this function will try to match the coordinates to a
        location using the Google Maps API.

        Returns a string: City, Region, Country
        """
        if ipaddr in self.ip_and_location:
            return self.ip_and_location[ipaddr]
        else:
            if self.IPINFO_TOKEN:
                # location_json = self.lookup_ip(event.details['browser']['address'])
                location_json = self.lookup_ip(ipaddr)
                if location_json:
                    city = region = country = "Unknown"
                    if "city" in location_json:
                        if location_json['city']:
                            city = location_json['city']
                    if "region" in location_json:
                        if location_json['region']:
                            region = location_json['region']
                    if "country" in location_json:
                        if location_json['country']:
                            country = location_json['country']
                    location = f"{city}, {region}, {country}"
                else:
                    location = f"{target.latitude}, {target.longitude}"
            elif google:
                if self.GEOLOCATE_TOKEN:
                    location = self.get_google_location_data(target.latitude, target.longitude)
                else:
                    location = f"{target.latitude}, {target.longitude}"
            else:
                location = f"{target.latitude}, {target.longitude}"
            self.locations.append(location)
            self.ip_and_location[ipaddr] = location
            return location

    def compare_ip_addresses(self, target_ip, browser_ip, verbose):
        """Compare the IP addresses of the target to that of an event. The goal: Looking for a
        mismatch that might identify some sort of interesting event. This might indicate an
        email was forwarded, a VPN was switched on/off, or maybe the target is at home.
        """
        if target_ip == browser_ip:
            return target_ip
        else:
            # We have an IP mismatch -- hard to tell why this might be.
            if verbose:
                print(f"[*] Event: This target's ({target_ip}) URL was clicked from a browser at {browser_ip}.")
            # This is an IP address not included in the results model, so we add it to our list here
            self.ip_addresses.append(browser_ip)
            return browser_ip

    def get_basic_campaign_info(self):
        """"Helper function to collect a campaign's basic details. This includes campaign name,
        status, template, and other details that are not the campaign's results.

        This keeps these calls in one place for tidiness and easier management.
        """
        self.cam_name = self.campaign.name
        self.cam_status = self.campaign.status
        self.created_date = self.campaign.created_date
        self.launch_date = self.campaign.launch_date
        self.completed_date = self.campaign.completed_date
        self.cam_url = self.campaign.url

        # Collect SMTP information
        self.smtp = self.campaign.smtp
        self.cam_from_address = self.smtp.from_address
        self.cam_smtp_host = self.smtp.host

        # Collect the template information
        self.template = self.campaign.template
        self.cam_subject_line = self.template.subject
        self.cam_template_name = self.template.name
        self.cam_template_attachments = self.template.attachments
        if self.cam_template_attachments == []:
            self.cam_template_attachments = "None Used"

        # Collect the landing page information
        self.page = self.campaign.page
        self.cam_page_name = self.page.name
        self.cam_redirect_url = self.page.redirect_url
        if self.cam_redirect_url == "":
            self.cam_redirect_url = "Not Used"
        self.cam_capturing_passwords = self.page.capture_passwords
        self.cam_capturing_credentials = self.page.capture_credentials

    def collect_all_campaign_info(self, combine_reports):
        """Collect the campaign's details and set values for each of the variables."""
        # Collect the basic campaign details
        try:
            # Begin by checking if the ID is valid
            self.cam_id = self.campaign.id
            if combine_reports and self.cam_name is None:
                print(f"[+] Reports will be combined -- setting name, dates, and URL based on campaign ID {self.cam_id}.")
                self.get_basic_campaign_info()
            elif combine_reports is False:
                self.get_basic_campaign_info()
            # Collect the results and timeline lists
            if self.results is None:
                self.results = self.campaign.results
                self.timeline = self.campaign.timeline
            elif combine_reports:
                self.results += self.campaign.results
                self.timeline += self.campaign.timeline
            else:
                self.results = self.campaign.results
                self.timeline = self.campaign.timeline
        except:
            print(f"[!] Looks like campaign ID {self.cam_id} does not exist! Skipping it...")

    def process_results(self, combine_reports):
        """Process the results model to collect basic data, like total targets and event details.
        This should be run after the process_timeline_events() function which creates the
        targets_* lists.

        The results model can provide:
        first_name, last_name, email, position, and IP address
        """
        # Total length of results gives us the total number of targets
        if combine_reports and self.total_targets is None:
            self.total_targets = len(self.campaign.results)
        elif combine_reports:
            self.total_targets += len(self.campaign.results)
        else:
            # Not combining, so reset counters
            self.total_unique_opened = 0
            self.total_unique_clicked = 0
            self.total_unique_reported = 0
            self.total_unique_submitted = 0
            # Reports will not be combined, so reset tracking between reports
            self.total_targets = len(self.campaign.results)
            self.ip_addresses = []
            self.campaign_results_summary = []
        # Go through all results and extract data for statistics
        for target in self.campaign.results:
            temp_dict = {}
            # Log the IP address for additional statistics later
            if not target.ip == "":
                self.ip_addresses.append(target.ip)
                self.geolocate(target, target.ip, self.google)
            # Add all of the recipient's details and results to the temp dictionary
            temp_dict["email"] = target.email
            temp_dict["fname"] = target.first_name
            temp_dict["lname"] = target.last_name
            position = "None Provided"
            if target.position:
                position = target.position
            temp_dict["position"] = position
            temp_dict["ip_address"] = target.ip
            # Check if this target was recorded as viewing the email (tracking image)
            if target.email in self.targets_opened:
                temp_dict["opened"] = True
                self.total_unique_opened += 1
            else:
                temp_dict["opened"] = False
            # Check if this target clicked the link
            if target.email in self.targets_clicked:
                temp_dict["clicked"] = True
                self.total_unique_clicked += 1
                # Incremement the total number of opens for this target if they clicked
                # but did not display the tracking image in the email
                if target.email not in self.targets_opened:
                    self.total_unique_opened += 1
            else:
                temp_dict["clicked"] = False
            # Check if this target submitted data
            if target.email in self.targets_submitted:
                temp_dict["submitted"] = True
                self.total_unique_submitted += 1
            else:
                temp_dict["submitted"] = False
            # Check if this target reported the email
            if target.email in self.targets_reported:
                temp_dict["reported"] = True
                self.total_unique_reported += 1
            else:
                temp_dict["reported"] = False
            # Append the temp dictionary to the event summary list
            self.campaign_results_summary.append(temp_dict)

    def process_timeline_events(self, combine_reports):
        """Process the timeline model to collect basic data, like total clicks, and get detailed
        event data for recipients.

        The timeline model contains all events that occurred during the campaign.
        """
        # Create counters for enumeration
        sent_counter = 0
        click_counter = 0
        opened_counter = 0
        reported_counter = 0
        submitted_counter = 0

        # Reset target lists
        self.targets_opened = []
        self.targets_clicked = []
        self.targets_reported = []
        self.targets_submitted = []
        # Run through all events and count each of the four basic events
        for event in self.campaign.timeline:
            if event.message == "Email Sent":
                sent_counter += 1
            elif event.message == "Email Opened":
                opened_counter += 1
                self.targets_opened.append(event.email)
            elif event.message == "Clicked Link":
                click_counter += 1
                self.targets_clicked.append(event.email)
            elif event.message == "Submitted Data":
                submitted_counter += 1
                self.targets_submitted.append(event.email)
            elif event.message == "Email Reported":
                reported_counter += 1
                self.targets_reported.append(event.email)
        # Assign the counter values to our tracking lists
        if combine_reports:
            # Append, +=, totals if combining reports
            self.total_sent += sent_counter
            self.total_opened += opened_counter
            self.total_clicked += click_counter
            self.total_reported += reported_counter
            self.total_submitted += submitted_counter
        else:
            # Set tracking variables to current counter values for non-combined reports
            self.total_sent = sent_counter
            self.total_opened = opened_counter
            self.total_clicked = click_counter
            self.total_reported = reported_counter
            self.total_submitted = submitted_counter

    def generate_report(self):
        """Determines which type of report generate and the calls the appropriate reporting
        functions.
        """
        if self.report_format == "excel":
            print("[+] Building the report -- you selected a Excel/xlsx report.")
            self.output_xlsx_report = self._build_output_xlsx_file_name()
            self.write_xlsx_report()
        elif self.report_format == "word":
            print("[+] Building the report -- you selected a Word/docx report.")
            print("[+] Looking for the a template_[...].docx-File to be used for the Word report.")

            # Check if a template-Files is in dictionary files (must be 3 files)
            onlyfiles = [f for f in listdir("./files") if isfile(join("./files", f))]
            count_templates  =0
            templates_are_in = False
            for file in onlyfiles:
                if re.match("template.*docx", file):
                    count_templates += 1
                    print(file)
                    if count_templates == 3:
                        templates_are_in = True
                        break

            if templates_are_in:
                print("[+] Required Templates was found -- proceeding with report generation...")
                print("L.. Word reports can take a while if you had a lot of recipients.")
                self.output_word_report = self._build_output_word_file_name()
                self.write_word_report()
            else:
                print("[!] Could not find the template document! Make sure 'template.docx' is in the GoReport directory.")
                sys.exit()
        elif self.report_format == "quick":
            print("[+] Quick report stats:")
            self.get_quick_stats()

    def get_quick_stats(self):
        """Present quick stats for the campaign. Just basic numbers and some details."""
        print()
        print(self.cam_name)
        print(f"Status:\t\t{self.cam_status}")
        print(f"Created:\t{self.created_date.split('T')[1].split('.')[0]} on {self.created_date.split('T')[0]}")
        print(f"Started:\t{self.launch_date.split('T')[1].split('.')[0]} on {self.launch_date.split('T')[0]}")
        if self.cam_status == "Completed":
            print(f"Completed:\t{self.completed_date.split('T')[1].split('.')[0]} on {self.completed_date.split('T')[0]}")
        print()
        print(f"Total Targets:\t{self.total_targets}")
        print(f"Emails Sent:\t{self.total_sent}")
        print(f"IPs Seen:\t{len(self.ip_addresses)}")
        print()
        print(f"Total Opened Events:\t\t{self.total_opened}")
        print(f"Total Click Events:\t\t{self.total_clicked}")
        print(f"Total Submitted Data Events:\t{self.total_submitted}")
        print()
        print(f"Individuals Who Opened:\t\t\t{self.total_unique_opened}")
        print(f"Individuals Who Clicked:\t\t{self.total_unique_clicked}")
        print(f"Individuals Who Entered Data:\t\t{self.total_unique_submitted}")
        print(f"Individuals Who Reported the Email:\t{self.total_unique_reported}")

    def _build_output_xlsx_file_name(self):
        """Create the xlsx report name."""
        safe_name = "".join([c for c in self.cam_name if c.isalpha() or c.isdigit() or c == " "]).rstrip()
        xlsx_report = f"Gophish Results for {safe_name}.xlsx"
        return xlsx_report

    def _build_output_word_file_name(self):
        """Create the docx report name."""
        safe_name = "".join([c for c in self.cam_name if c.isalpha() or c.isdigit() or c == " "]).rstrip()
        word_report = f"Gophish Results for {safe_name}_{self.company}.docx"
        return word_report

    def _set_word_column_width(self, column, width):
        """Custom function for quickly and easily setting the width of a table's column in the Word
        docx output.

        This option is missing from the basic Python-docx library.
        """
        for cell in column.cells:
            cell.width = width

    def write_xlsx_report(self):
        """Assemble and output the xlsx file report.

        Throughout this function, results are assembled by adding commas and then adding to a
        results string, i.e. 'result_A' and then 'result_A' += ',result_B'. This is so the
        result can be written to the csv file and have the different pieces end up in the correct
        columns.
        """
        goreport_xlsx = xlsxwriter.Workbook(self.output_xlsx_report)
        # Bold format
        bold_format = goreport_xlsx.add_format({'bold': True})
        bold_format.set_text_wrap()
        bold_format.set_align('vcenter')
        # Centered format
        center_format = goreport_xlsx.add_format()
        center_format.set_text_wrap()
        center_format.set_align('vcenter')
        center_format.set_align('center')
        # Header format
        header_format = goreport_xlsx.add_format({'bold': True})
        header_format.set_text_wrap()
        header_format.set_align('vcenter')
        header_format.set_bg_color(self.xlsx_header_bg_color)
        header_format.set_font_color(self.xlsx_header_font_color)
        # Number cells
        num_format = goreport_xlsx.add_format()
        num_format.set_align('center')
        # Boolean cells - True
        true_format = goreport_xlsx.add_format({'bold': True})
        true_format.set_text_wrap()
        true_format.set_align('vcenter')
        true_format.set_font_color("#9C0006")
        true_format.set_bg_color("#FFC7CE")
        # Boolean cells - True
        false_format = goreport_xlsx.add_format()
        false_format.set_text_wrap()
        false_format.set_align('vcenter')
        false_format.set_font_color("#006100")
        false_format.set_bg_color("#C6EFCE")
        # Remaining cells
        wrap_format = goreport_xlsx.add_format()
        wrap_format.set_text_wrap()
        wrap_format.set_align('vcenter')

        worksheet = goreport_xlsx.add_worksheet("Overview")
        col = 0
        row = 0

        worksheet.set_column(0, 10, 62)

        worksheet.write(row, col, "Campaign Results For:", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_name}", wrap_format)
        row += 1
        worksheet.write(row, col, "Status", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_status}", wrap_format)
        row += 1
        worksheet.write(row, col, "Created", bold_format)
        worksheet.write(row, col + 1, f"{self.created_date}", wrap_format)
        row += 1
        worksheet.write(row, col, "Started", bold_format)
        worksheet.write(row, col + 1, f"{self.launch_date}", wrap_format)
        row += 1
        if self.cam_status == "Completed":
            worksheet.write(row, col, "Completed", bold_format)
            worksheet.write(row, col + 1, f"{self.completed_date}", wrap_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Campaign Details", bold_format)
        row += 1
        worksheet.write(row, col, "From", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_from_address}", wrap_format)
        row += 1
        worksheet.write(row, col, "Subject", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_subject_line}", wrap_format)
        row += 1
        worksheet.write(row, col, "Phish URL", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_url}", wrap_format)
        row += 1
        worksheet.write(row, col, "Redirect URL", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_redirect_url}", wrap_format)
        row += 1
        worksheet.write(row, col, "Attachment(s)", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_template_attachments}", wrap_format)
        row += 1
        worksheet.write(row, col, "Captured Passwords", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_capturing_credentials}", wrap_format)
        row += 1
        worksheet.write(row, col, "Stored Passwords", bold_format)
        worksheet.write(row, col + 1, f"{self.cam_capturing_passwords}", wrap_format)
        row += 1

        worksheet.write(row, col, "")
        row += 1

        # Write a high level summary for stats
        worksheet.write(row, col, "High Level Results", bold_format)
        row += 1
        worksheet.write(row, col, "Total Targets", bold_format)
        worksheet.write(row, col + 1, self.total_targets, num_format)
        row += 1

        worksheet.write(row, col, "The following totals indicate how many events of each type Gophish recorded:", wrap_format)
        row += 1
        worksheet.write(row, col, "Total Opened Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_opened, num_format)
        row += 1
        worksheet.write(row, col, "Total Clicked Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_clicked, num_format)
        row += 1
        worksheet.write(row, col, "Total Submitted Data Events", bold_format)
        worksheet.write(row, col + 1, "", wrap_format)
        row += 1
        worksheet.write(row, col, "Total Report Events", bold_format)
        worksheet.write_number(row, col + 1, self.total_reported, num_format)
        row += 1

        worksheet.write(row, col, "The following totals indicate how many targets participated in each event type:", wrap_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Opened", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_opened, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Clicked", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_clicked, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Submitted Data", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_submitted, num_format)
        row += 1
        worksheet.write(row, col, "Individuals Who Reported", bold_format)
        worksheet.write_number(row, col + 1, self.total_unique_reported, num_format)
        row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet = goreport_xlsx.add_worksheet("Summary")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 20)

        worksheet.write(row, col, "Summary of Events", bold_format)
        row += 1

        header_col = 0
        headers = ["Email Address", "Open", "Click", "Creds", "Report", "OS", "Browser"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1

        # Sort campaign summary by each dict's email entry and then create results table
        target_counter = 0
        ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])               # Kampanien-resulate-zusammenfassung sortiert nach email
        for target in ordered_results:
            worksheet.write(row, col, target['email'], wrap_format)
            if target['opened']:
                worksheet.write_boolean(row, col + 1, target['opened'], true_format)
            else:
                worksheet.write_boolean(row, col + 1, target['opened'], false_format)
            if target['clicked']:
                worksheet.write_boolean(row, col + 2, target['clicked'], true_format)
            else:
                worksheet.write_boolean(row, col + 2, target['clicked'], false_format)
            if target['submitted']:
                worksheet.write_boolean(row, col + 3, target['submitted'], true_format)
            else:
                worksheet.write_boolean(row, col + 3, target['submitted'], false_format)
            if target['reported']:
                worksheet.write_boolean(row, col + 4, target['reported'], true_format)
            else:
                worksheet.write_boolean(row, col + 4, target['reported'], false_format)
            if target['email'] in self.targets_clicked:
                for event in self.timeline:
                    if event.message == "Clicked Link" and event.email == target['email']:
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 5, browser_details, wrap_format)
                        worksheet.write(row, col + 6, os_details, wrap_format)
            else:
                worksheet.write(row, col + 5, "N/A", wrap_format)
                worksheet.write(row, col + 6, "N/A", wrap_format)
            row += 1
            target_counter += 1
            print(f"[+] Created row for {target_counter} of {self.total_targets}.")

        print("[+] Finished writing events summary...")
        print("[+] Detailed results analysis is next and will take some time if you had a lot of targets...")
        # End of the event summary and beginning of the detailed results

        worksheet = goreport_xlsx.add_worksheet("Event Details")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 40)

        worksheet.write(row, col, "Detailed Analysis", bold_format)
        row += 1

        target_counter = 0
        for target in self.results:
            # Only create a Detailed Analysis section for targets with clicks
            if target.email in self.targets_clicked:
                position = ""
                if target.position:
                    position = f"({target.position})"
                worksheet.write(row, col, f"{target.first_name} {target.last_name} {position}", bold_format)
                row += 1
                worksheet.write(row, col, target.email, wrap_format)
                row += 1
                # Go through all events to find events for this target
                for event in self.timeline:
                    if event.message == "Email Sent" and event.email == target.email:
                        # Parse the timestamp into separate date and time variables
                        temp = event.time.split('T')
                        sent_date = temp[0]
                        sent_time = temp[1].split('.')[0]
                        # Record the email sent date and time in the report
                        worksheet.write(row, col, f"Sent on {sent_date.replace(',', '')} at {sent_time}", wrap_format)
                        row += 1

                    if event.message == "Email Opened" and event.email == target.email:
                        # Record the email preview date and time in the report
                        temp = event.time.split('T')
                        worksheet.write(row, col, f"Email Preview at {temp[0]} {temp[1].split('.')[0]}", wrap_format)
                        row += 1

                    if event.message == "Clicked Link" and event.email == target.email:
                        worksheet.write(row, col, "Email Link Clicked", bold_format)
                        row += 1

                        header_col = 0
                        headers = ["Time", "IP", "Location", "Browser", "Operating System"]
                        for header in headers:
                            worksheet.write(row, header_col, header, header_format)
                            header_col += 1
                        row += 1

                        temp = event.time.split('T')
                        worksheet.write(row, col, f"{temp[0]} {temp[1].split('.')[0]}", wrap_format)

                        # Check if browser IP matches the target's IP and record result
                        ip_comparison = self.compare_ip_addresses(target.ip,
                                                                  event.details['browser']['address'],
                                                                  self.verbose)
                        worksheet.write(row, col + 1, f"{ip_comparison}", wrap_format)

                        # Parse the location data
                        loc = self.geolocate(target, event.details['browser']['address'], self.google)
                        worksheet.write(row, col + 2, loc, wrap_format)

                        # Parse the user-agent string and add browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        worksheet.write(row, col + 3, browser_details, wrap_format)
                        self.browsers.append(browser_details)

                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 4, os_details, wrap_format)
                        self.operating_systems.append(os_details)
                        row += 1

                    if event.message == "Submitted Data" and event.email == target.email:
                        # Now we have events for submitted data. A few notes on this:
                        #   1. There is no expectation of a Submit event without a Clicked Link event
                        #   2. Assuming that, the following process does NOT flag IP mismatches
                        #      or add to the list of seen locations, OSs, IPs, or browsers.
                        worksheet.write(row, col, "Submitted Data Captured", bold_format)
                        row += 1

                        header_col = 0
                        headers = ["Time", "IP", "Location", "Browser", "Operating System", "Data Captured"]
                        for header in headers:
                            worksheet.write(row, header_col, header, header_format)
                            header_col += 1
                        row += 1

                        temp = event.time.split('T')
                        worksheet.write(row, col, f"{temp[0]} {temp[1].split('.')[0]}", wrap_format)

                        worksheet.write(row, col + 1, f"{event.details['browser']['address']}", wrap_format)

                        loc = self.geolocate(target, event.details['browser']['address'], self.google)
                        worksheet.write(row, col + 2, loc, wrap_format)

                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        worksheet.write(row, col + 3, browser_details, wrap_format)

                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        worksheet.write(row, col + 4, os_details, wrap_format)

                        # Get just the submitted data from the event's payload
                        submitted_data = ""
                        data_payload = event.details['payload']
                        # Get all of the submitted data
                        for key, value in data_payload.items():
                            # To get just submitted data, we drop the 'rid' key
                            if not key == "rid":
                                submitted_data += f"{key}:{str(value).strip('[').strip(']')}"
                        worksheet.write(row, col + 5, submitted_data, wrap_format)
                        row += 1

                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
            else:
                # This target had no clicked or submitted events so move on to next
                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
                continue
            worksheet.write(row, col, "")
            row += 1

        print("[+] Finished writing detailed analysis...")

        worksheet = goreport_xlsx.add_worksheet("Stats")
        row = 0
        col = 0

        worksheet.set_column(0, 10, 35)

        worksheet.write(row, col, "Recorded Browsers Based on User-Agents:", bold_format)
        row += 1

        header_col = 0
        headers = ["Browser", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_browsers = Counter(self.browsers)
        for key, value in counted_browsers.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Record OS From Browser User-Agents:", bold_format)
        row += 1
        header_col = 0
        headers = ["Operating System", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_os = Counter(self.operating_systems)
        for key, value in counted_os.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Recorded Locations from IPs:", bold_format)
        row += 1
        header_col = 0
        headers = ["Locations", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_locations = Counter(self.locations)
        for key, value in counted_locations.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "")
        row += 1

        worksheet.write(row, col, "Recorded IPs:", bold_format)
        row += 1
        header_col = 0
        headers = ["IP Address", "Seen"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        counted_ip_addresses = Counter(self.ip_addresses)
        for key, value in counted_ip_addresses.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write_number(row, col + 1, value, num_format)
            row += 1

        worksheet.write(row, col, "Recorded IPs and Locations:", bold_format)
        row += 1
        header_col = 0
        headers = ["IP Address", "Location"]
        for header in headers:
            worksheet.write(row, header_col, header, header_format)
            header_col += 1
        row += 1
        for key, value in self.ip_and_location.items():
            worksheet.write(row, col, f"{key}", wrap_format)
            worksheet.write(row, col + 1, f"{value}", wrap_format)
            row += 1

        goreport_xlsx.close()
        print(f"[+] Done! Check '{self.output_xlsx_report}' for your results.")

    def write_word_report(self):
        """Assemble and output the Word docx file report."""
        # Create document writer using the template and a style editor

        if self.company == 'Fraunhofer':
            d = Document("./files/template_Fraunhofer.docx")
            underline = './files/FH_underline.png'
            table_style = 'Fraunhofer'
        elif self.company == 'HSMW':
            d = Document("./files/template_HSMW.docx")
            underline = './files/HSMW_underline.png'
            table_style = 'HSMW'
        elif self.company == 'Axilaris':
            d = Document("./files/template_Axilaris.docx")
            underline = './files/Axilaris_underline.png'
            table_style = 'Axilaris'
        styles = d.styles

        styles['Heading 1'].font.color.rgb = docx.shared.RGBColor(44, 62, 80)  # HEading Farbe ändern
        styles['Heading 1'].font.size = Pt(28)
        styles['Heading 1'].font.name = 'Calibri'
        styles['Heading 1'].font.bold = False

        # Create a custom styles for table cells
        _ = styles.add_style("Cell Text", WD_STYLE_TYPE.CHARACTER)
        cell_text = d.styles["Cell Text"]
        cell_text_font = cell_text.font
        cell_text_font.name = "Calibri"
        cell_text_font.size = Pt(12)
        cell_text_font.bold = True
        cell_text_font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        _ = styles.add_style("Cell Text Hit", WD_STYLE_TYPE.CHARACTER)
        cell_text_hit = d.styles["Cell Text Hit"]
        cell_text_hit_font = cell_text_hit.font
        cell_text_hit_font.name = "Calibri"
        cell_text_hit_font.size = Pt(12)
        cell_text_hit_font.bold = True
        cell_text_hit_font.color.rgb = RGBColor(0x00, 0x96, 0x00)

        _ = styles.add_style("Cell Text Miss", WD_STYLE_TYPE.CHARACTER)
        cell_text_miss = d.styles["Cell Text Miss"]
        cell_text_miss_font = cell_text_miss.font
        cell_text_miss_font.name = "Calibri"
        cell_text_miss_font.size = Pt(12)
        cell_text_miss_font.bold = True
        cell_text_miss_font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        def add_heading(title):
            d.add_heading(f'{title}')
            p = d.add_paragraph()
            p.add_run().add_picture(underline)
            p.add_run().add_break()

        def create_fig(labels, data, title):

            fig, ax = plt.subplots(figsize=(3, 4.5), subplot_kw=dict(aspect="equal"))  # Definition eines Plot Objekts
            plt.subplots_adjust(bottom=0.3)  # Definition wie sehr Titel,Chart und Legende zusammengerückt werden

            labels = labels  # Definition der Label
            data = data  # Definition der Daten


            # Definition der Farbvarianten für unterschiedliche Fälle
            if len(labels) == 2:
                colors = ['#922B21', '#1F618D']
            elif len(labels) == 3:
                colors = ['#922B21','#1F618D', '#117A65']
            elif len(labels) == 4:
                colors = ['#922B21','#1F618D', '#117A65', '#B7950B']
            elif len(labels) == 5:
                colors = ['#922B21','#1F618D', '#117A65', '#B7950B','#A04000']
            elif len(labels) == 6:
                colors = ['#922B21','#1F618D', '#117A65', '#B7950B','#A04000','#717D7E']
            elif len(labels) == 7:
                colors = ['#922B21','#1F618D', '#117A65', '#B7950B','#A04000','#717D7E','#2E4053']
            elif len(labels) == 8:
                colors = ['#922B21','#1F618D', '#117A65', '#B7950B','#A04000','#717D7E','#2E4053','#C39BD3']
            elif len(labels) == 9:
                colors = ['#922B21','#1F618D', '#117A65', '#B7950B','#A04000','#717D7E','#2E4053','#C39BD3','#A3E4D7']
            elif len(labels) == 10:
                colors = ['#922B21','#1F618D', '#117A65', '#B7950B','#A04000','#717D7E','#2E4053','#C39BD3','#A3E4D7','#FAD7A0']


            # Initialisierung Pie Chart
            plt.pie(data, labels=data, colors=colors,
                    textprops={'fontsize': 11},  # Definition der Schriftgröße
                    startangle=90,
                    wedgeprops={'edgecolor': 'Black',  # Definition des Rahmen einer Sektion
                                'linewidth': 0.9,
                                'antialiased': True})

            ax.axis('equal')  # Auto Anpassung an Bildgröße

            ncol = 0
            if len(labels) > 5:  # dass Legende nicht zu groß wird
                ncol = 2
            else:
                ncol = 1

            # Initialisierung der Legende
            ax.legend(labels=labels, loc='lower center', bbox_to_anchor=(0.5, -0.44)
                      # Posotionierung der Legende mit Farbanpassung und Größe
                      , fontsize=11, edgecolor='Black', facecolor='#E8E8E8',
                      ncol=ncol)

            # Initialisierung des Titels
            ax.set_title(f"{title}:", loc='center', y=1,
                         # Posotionierung des Titels mit Farbanpassung und Größe
                         fontdict={'fontsize': 16})

            plt.savefig(f'./files/{title}.png', bbox_inches='tight', transparent=True)  # Abspeichern in Datei

        def time_chart():

            sns.set(rc={'figure.figsize': (10, 8)})
            sns.set_style("darkgrid", {"grid.color": ".6", "grid.linestyle": ":"})

            ax = sns.pointplot(data=self.time_df, x='time', y='count',
                               hue="event",
                               hue_order=["opened", "clicked", "submitted", "reported"],
                               palette='colorblind', linestyles="dotted")

            # legend='full', lw=3)

            # set Title
            ax.set_title("Event Übersicht", fontsize=24)
            # Edit X-Axis Propertiy
            plt.setp(ax.xaxis.get_majorticklabels(), rotation=90)
            ax.xaxis.set_minor_locator(md.DayLocator(interval=2))
            ax.tick_params(axis='x', which='major', length=10)
            ax.tick_params(axis='x', which='minor', length=5)
            # Set Legend Size
            plt.legend(bbox_to_anchor=(1, 1), prop={"size": 15})
            # Set Label
            ax.set_ylabel("Häufigkeit der Events", fontsize=20)
            ax.set_xlabel("Zeitraum", fontsize=20)

            plt.savefig('./files/timeline.png', bbox_inches='tight', transparent=True)

        def portscan(ip_addr):
            scanner = nmap.PortScanner()

            d.add_page_break()
            p = d.add_paragraph()

            version_num = re.findall("\d+",str(scanner.nmap_version()))

            d.add_heading(f'Nmap-Portscan (version: {version_num[0]}.{version_num[1]})')


            scan_dict = scanner.scan(ip_addr, "1-1024", arguments='-v -sS -sV -sC -A -O')

            table = d.add_table(2, 2)

            if scanner[ip_addr].state() == 'up':
                print("[+] Port Scan erfolgreich!")
                cell_1 = table.cell(0, 0)
                cell_1.text = f'''
IP-Adresse: {scan_dict['scan'][ip_addr]['addresses']['ipv4']}
Hostname: {scan_dict['scan'][ip_addr]['hostnames'][0]['name']} '''

                cell_2 = table.cell(0, 1)
                cell_2.text = f'''
Scanner Status: {scanner[ip_addr].state()}
Zeitstempel: {datetime.strptime(scan_dict['nmap']['scanstats']['timestr'], '%c')} '''

                cell_3 = table.cell(1, 0)
                cell_3.text = ''
                cell_3.paragraphs[0].add_run("Operating Systems: ")

                table_2 = d.add_table((len(list(scan_dict['scan'][ip_addr]['tcp'].keys()))), 1, style="port_tab")

                for i in range(len(scanner[ip_addr]['osmatch']) - 1):
                    cell_3.paragraphs[0].add_run(f'''
- {scanner[ip_addr]['osmatch'][i]['name']} ''')

                for k, v in (scan_dict['scan'][ip_addr]['tcp'].items()):

                    cpe, info, state, version, name = v['cpe'], v['extrainfo'], v['state'], v['version'], v['name']
                    try:
                        script = v['script']
                    except KeyError:
                        script = ''

                    cell = table_2.cell(int(list(scan_dict['scan'][ip_addr]['tcp'].keys()).index(k)), 0)
                    cell.text = ''
                    cell.paragraphs[0].add_run(f'Port: {k}').bold = True
                    cell.paragraphs[0].add_run(f'''
    name: {name}
    cpe: {cpe}
    info: {info}
    scriptinfo: {script}
    state: {state}
    version: {version} ''')




            else:
                print("[+] Port Scan nicht erfolgreich!")
                p.add_run(f"""
Scanner Status: {scanner[ip_addr].state()}
Portscan war nicht erfolgreich. """)

        # HInzufuegen einer Kampagnen-Zusammenfassung

        print(self.cam_status)


        add_heading('Zusammenfassung')
        table = d.add_table(2, 2)
        cell_1 = table.cell(0, 0)
        cell_1.text = ''
        cell_1.paragraphs[0].add_run(f"Kampagne Resultate für: {self.cam_name}").bold = True
        # Runs are basically "runs" of text and must be aligned like we want
        # them aligned in the report -- thus they are pushed left
        if self.cam_status == "Completed":
            completed_status = f"Abgeschlossen am:\t{self.completed_date.split('T')[1].split('.')[0]} um {self.completed_date.split('T')[0]}"
        else:
            completed_status = "Noch Aktiv"

        if self.cam_status == 'In progress':
            status = 'Im Gange'
        else:
            status = self.cam_status
        cell_1.paragraphs[0].add_run(f'''
Status: {status}
Erstellt: {self.created_date.split('T')[1].split('.')[0]} am {self.created_date.split('T')[0]}
Gestartet: {self.launch_date.split('T')[1].split('.')[0]} am {self.launch_date.split('T')[0]}
Abgeschlossen: {completed_status}
''')

        if self.cam_status == "Completed":
            print()

        # HInzufuegen einer Kampagnen-Details
        cell_2 = table.cell(0, 1)
        cell_2.text = ''
        cell_2.paragraphs[0].add_run("Kampagnendetails:").bold = True

        # Eindeutschen
        if self.cam_capturing_credentials == 'False':
            capturing_credentials = 'Falsch'
        elif self.cam_capturing_credentials == 'True':
            capturing_credentials = 'Wahr'
        else:
            capturing_credentials = self.cam_capturing_credentials

        if self.cam_capturing_passwords == 'False':
            capturing_passwords = 'Falsch'
        elif self.cam_capturing_passwords == 'True':
            capturing_passwords = 'Wahr'
        else:
            capturing_passwords = self.cam_capturing_passwords

        if self.cam_template_attachments == 'None Used':
            template_attachments = 'Nicht Verwendet'
        elif self.cam_template_attachments == 'Used':
            template_attachments = 'Verwendet'
        else:
            template_attachments = self.cam_template_attachments

        if self.cam_redirect_url == 'Not Used':
            redirect_url = 'Nicht Benutzt'
        elif self.cam_redirect_url == 'Used':
            redirect_url = 'Benutzt'
        else:
            redirect_url = self.cam_redirect_url




        cell_2.paragraphs[0].add_run(f"""
Von: {self.cam_from_address}
Thema: {self.cam_subject_line}
Phishing-URL: {self.cam_url}
Umleitungs-URL: {redirect_url}
Anhänge: {template_attachments}
Erfasste Anmeldedaten: {capturing_credentials}
Gespeicherte Passwörter: {capturing_passwords}
""")
        d.add_page_break()

        # Write a high level summary for stats
        add_heading('Zusammenfassung der Events')


        event_labels = ['Geöffnet', 'Angeklicked', 'Gemeldet', 'Eingesendet']
        event_type_data = [self.total_opened,self.total_clicked,self.total_reported,self.total_submitted]
        targets_in_type_data = [self.total_unique_opened,self.total_unique_clicked,self.total_unique_reported,self.total_unique_submitted]

        create_fig(event_labels,event_type_data,'Event Typen')
        create_fig(event_labels,targets_in_type_data,'Beteiligung an Events')

        p = d.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(f'Total Targets: {self.total_targets}')
        run.bold = True
        run.font.size = Pt(14)

        table = d.add_table(2, 2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        temp_cell = table.cell(1, 0)
        temp_cell.text = 'Summen, wie viele Event Typen aufgezeichnet wurden'
        temp_cell = table.cell(0, 0)
        temp_cell.paragraphs[0].add_run().add_picture('./files/Event Typen.png')
        temp_cell = table.cell(1, 1)
        temp_cell.text = 'Summen, wie viele Ziele an jedem Event Typ beteiligt waren'
        temp_cell = table.cell(0, 1)
        temp_cell.paragraphs[0].add_run().add_picture('./files/Beteiligung an Events.png')

        os.remove('./files/Event Typen.png')
        os.remove('./files/Beteiligung an Events.png')

        d.add_page_break()

        print("[+] Finished writing high level summary...")
        # Beginn der Event-Zusammenfassung
        add_heading("Detailierte Zusammenfassung der Events")
        d.add_paragraph("Die folgende Tabelle fasst zusammen, wer die in dieser Kampagne gesendeten E-Mails geöffnet und angeklickt hat.")

        # Erstellen einer Tabelle für die Event Zusammenfassung
        table = d.add_table(rows=len(self.campaign_results_summary) + 1, cols=7, style=table_style)
        #Initialisierung für Header in Tabelle
        headerlist = ["Email Adresse", "Geöffnet", "Angeklickt", "Erfasst", "Gemeldet", "OS", "Browser"]
        for i in range(7):
            header = f'header{i}'
            header = table.cell(0, i)
            header.text = ""
            header.paragraphs[0].add_run(headerlist[i], "Cell Text").bold = True

        # Sortieren der Einträge nach Email
        target_counter = 0
        counter = 1
        ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])
        for target in ordered_results:
            email_cell = table.cell(counter, 0)
            email_cell.text = f"{target['email']}"

            temp_cell = table.cell(counter, 1)
            if target['opened']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 2)
            if target['clicked']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 3)
            if target['submitted']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 4)
            if target['reported']:
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            if target['email'] in self.targets_clicked:
                for event in self.timeline:
                    if event.message == "Clicked Link" and event.email == target['email']:
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string

                        os_details = user_agent.os.family + " " + \
                            user_agent.os.version_string
                        temp_cell = table.cell(counter, 5)
                        temp_cell.text = os_details
                        temp_cell = table.cell(counter, 6)
                        temp_cell.text = browser_details
            else:
                temp_cell = table.cell(counter, 5)
                temp_cell.text = "N/A"
                temp_cell = table.cell(counter, 6)
                temp_cell.text = "N/A"
            counter += 1
            target_counter += 1
            print(f"[+] Created table entry for {target_counter} of {self.total_targets}.")

        d.add_page_break()

        # End of the event summary and beginning of the detailed results
        print("[+] Finished writing events summary...")
        print("[+] Detailed results analysis is next and may take some time if you had a lot of targets...")
        add_heading("Detaillierte Ergebnisse")
        target_counter = 0
        for target in self.results:
            # Only create a Detailed Analysis section for targets with clicks
            if target.email in self.targets_clicked:
                # Create counters to track table cell locations
                opened_counter = 1
                clicked_counter = 1
                submitted_counter = 1
                # Create section starting with a header with the first and last name
                position = ""
                if target.position:
                    position = f"({target.position})"
                d.add_heading(f"{target.first_name} {target.last_name} {position}", 2)
                p = d.add_paragraph(target.email)
                p = d.add_paragraph()
                # Save a spot to record the email sent date and time in the report
                email_sent_run = p.add_run()
                # Go through all events to find events for this target
                for event in self.timeline:
                    if event.message == "Email Sent" and event.email == target.email:
                        # Parse the timestamp into separate date and time variables
                        # Ex: 2017-01-30T14:31:22.534880731-05:00
                        temp = event.time.split('T')
                        sent_date = temp[0]
                        sent_time = temp[1].split('.')[0]

                        time = temp[1].split('.')[0]
                        date = temp[0]
                        date_time = temp[0] + " " + temp[1].split('.')[0]

                        if date in list(self.time_df["time"]):
                            try:
                                index = \
                                self.time_df.index[(self.time_df['time'] == date) & (self.time_df['event'] == "sent")].tolist()[0]
                            except:
                                index = (self.time_df.index[self.time_df["time"] == date].tolist())[0]
                            if self.time_df.loc[index, "time"] == date and self.time_df.loc[index, "event"] == "sent":
                                self.time_df.loc[index, "count"] += 1
                            else:
                                self.time_df.loc[len(list(self.time_df["time"]))] = date, 1, "sent"

                        else:
                            self.time_df.loc[len(list(self.time_df["time"]))] = date, 1, "sent"

                        # Record the email sent date and time in the run created earlier
                        email_sent_run.text = f"Email gesendet am {sent_date} um {sent_time} Uhr"

                    if event.message == "Email Opened" and event.email == target.email:
                        if opened_counter == 1:
                            # Create the Email Opened/Previewed table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("Email Geöffnet")
                            run.bold = True

                            opened_table = d.add_table(rows=1, cols=1, style=table_style)
                            opened_table.autofit = True
                            opened_table.allow_autofit = True

                            header1 = opened_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Zeit", "Cell Text").bold = True

                        # Begin by adding a row to the table and inserting timestamp
                        opened_table.add_row()
                        timestamp = opened_table.cell(opened_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        time = temp[1].split('.')[0]
                        date = temp[0]
                        date_time = temp[0] + " " + temp[1].split('.')[0]

                        if date in list(self.time_df["time"]):
                            try:
                                index = \
                                    self.time_df.index[
                                        (self.time_df['time'] == date) & (self.time_df['event'] == "opened")].tolist()[0]
                            except:
                                index = (self.time_df.index[self.time_df["time"] == date].tolist())[0]
                            if self.time_df.loc[index, "time"] == date and self.time_df.loc[index, "event"] == "opened":
                                self.time_df.loc[index, "count"] += 1
                            else:
                                self.time_df.loc[len(list(self.time_df["time"]))] = date, 1, "opened"

                        else:
                            self.time_df.loc[len(list(self.time_df["time"]))] = date, 1, "opened"

                        opened_counter += 1

                    if event.message == "Clicked Link" and event.email == target.email:
                        if clicked_counter == 1:
                            # Create the Clicked Link table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("Email-Link Angeklickt")
                            run.bold = True

                            clicked_table = d.add_table(rows=1, cols=5, style=table_style)
                            clicked_table.autofit = True
                            clicked_table.allow_autofit = True

                            headerlist = ["Zeit", "IP", "Position", "Browser", "OS"]
                            for i in range(5):
                                header = f'header{i}'
                                header = clicked_table.cell(0, i)
                                header.text = ""
                                header.paragraphs[0].add_run(headerlist[i], "Cell Text").bold = True



                        clicked_table.add_row()
                        timestamp = clicked_table.cell(clicked_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        time = temp[1].split('.')[0]
                        date = temp[0]
                        date_time = temp[0] + " " + temp[1].split('.')[0]

                        if date in list(self.time_df["time"]):
                            try:
                                index = \
                                    self.time_df.index[
                                        (self.time_df['time'] == date) & (self.time_df['event'] == "clicked")].tolist()[
                                        0]
                            except:
                                index = (self.time_df.index[self.time_df["time"] == date].tolist())[0]
                            if self.time_df.loc[index, "time"] == date and self.time_df.loc[index, "event"] == "clicked":
                                self.time_df.loc[index, "count"] += 1
                            else:
                                self.time_df.loc[len(list(self.time_df["time"]))] = date, 1, "clicked"

                        else:
                            self.time_df.loc[len(list(self.time_df["time"]))] = date, 1, "clicked"

                        ip_add = clicked_table.cell(clicked_counter, 1)
                        # Check if browser IP matches the target's IP and record result
                        ip_add.text = self.compare_ip_addresses(
                            target.ip, event.details['browser']['address'], self.verbose)


                        # Parse the location data
                        event_location = clicked_table.cell(clicked_counter, 2)
                        event_location.text = self.geolocate(target, event.details['browser']['address'], self.google)

                        # Parse the user-agent string for browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        browser = clicked_table.cell(clicked_counter, 3)
                        browser.text = browser_details
                        self.browsers.append(browser_details)
                        self.browsers_family.append(user_agent.browser.family)

                        op_sys = clicked_table.cell(clicked_counter, 4)
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        op_sys.text = os_details
                        self.operating_systems.append(os_details)
                        self.operating_systems_family.append(user_agent.os.family)

                        clicked_counter += 1




                    if event.message == "Submitted Data" and event.email == target.email:
                        if submitted_counter == 1:
                            # Create the Submitted Data table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("Daten Erfasst")
                            run.bold = True

                            submitted_table = d.add_table(rows=1, cols=6, style=table_style)
                            submitted_table.autofit = True
                            submitted_table.allow_autofit = True

                            headerlist = ["Zeit", "IP", "Position", "Browser", "OS", 'Daten Erfasst']
                            for i in range(6):
                                header = f'header{i}'
                                header = submitted_table.cell(0, i)
                                header.text = ""
                                header.paragraphs[0].add_run(headerlist[i], "Cell Text").bold = True



                        submitted_table.add_row()
                        timestamp = submitted_table.cell(submitted_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        time = temp[1].split('.')[0]
                        date = temp[0]
                        date_time = temp[0] + " " + temp[1].split('.')[0]

                        if date in list(self.time_df["time"]):
                            try:
                                index = \
                                    self.time_df.index[
                                        (self.time_df['time'] == date) & (self.time_df['event'] == "submitted")].tolist()[
                                        0]
                            except:
                                index = (self.time_df.index[self.time_df["time"] == date].tolist())[0]
                            if self.time_df.loc[index, "time"] == date and self.time_df.loc[index, "event"] == "submitted":
                                self.time_df.loc[index, "count"] += 1
                            else:
                                self.time_df.loc[len(list(self.time_df["time"]))] = date, 1, "submitted"

                        else:
                            self.time_df.loc[len(list(self.time_df["time"]))] = date, 1, "submitted"

                        ip_add = submitted_table.cell(submitted_counter, 1)
                        ip_add.text = event.details['browser']['address']

                        # Parse the location data
                        event_location = submitted_table.cell(submitted_counter, 2)
                        event_location.text = self.geolocate(target, event.details['browser']['address'], self.google)

                        # Parse the user-agent string and add browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])
                        browser_details = user_agent.browser.family + " " + \
                            user_agent.browser.version_string
                        browser = submitted_table.cell(submitted_counter, 3)
                        browser.text = browser_details

                        op_sys = submitted_table.cell(submitted_counter, 4)
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        op_sys.text = f"{os_details}"

                        # Get just the submitted data from the event's payload
                        submitted_data = ""
                        data = submitted_table.cell(submitted_counter, 5)
                        data_payload = event.details['payload']
                        # Get all of the submitted data
                        for key, value in data_payload.items():
                            # To get just submitted data, we drop the 'rid' key
                            if not key == "rid":
                                submitted_data += f"{key}:{str(value).strip('[').strip(']')}   "
                        data.text = f"{submitted_data}"
                        submitted_counter += 1

                        # Portscan if valid ip
                        ip_addr = self.compare_ip_addresses(
                            target.ip, event.details['browser']['address'], self.verbose)

                        # Proof if valid IP-Adress
                        if re.match(r'((25[0-5]|(2[0-4]|1\d|[1-9]|)\d)\.?\b){4}$', ip_addr):
                            print(f"[+] Start Port Scanning: {ip_addr}")
                            portscan(ip_addr)



                    if event.message == "Email Reported" and event.email == target.email:
                        temp = event.time.split('T')

                        time = temp[1].split('.')[0]
                        date = temp[0]
                        date_time = temp[0] + " " + temp[1].split('.')[0]        # Format ['2022-08-09', '13:50:49.6777583Z']   2022-08-09 13:50:49

                        if date in list(self.time_df["time"]):
                            try:
                                index = \
                                    self.time_df.index[
                                        (self.time_df['time'] == date) & (self.time_df['event'] == "reported")].tolist()[
                                        0]
                            except:
                                index = (self.time_df.index[self.time_df["time"] == date].tolist())[0]
                            if self.time_df.loc[index, "time"] == date and self.time_df.loc[index, "event"] == "reported":
                                self.time_df.loc[index, "count"] += 1
                            else:
                                self.time_df.loc[len(list(self.time_df["time"]))] = date, 1, "reported"

                        else:
                            self.time_df.loc[len(list(self.time_df["time"]))] = date, 1, "reported"

                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")

                d.add_page_break()
            else:
                # This target had no clicked or submitted events so move on to next
                target_counter += 1
                print(f"[+] Processed detailed analysis for {target_counter} of {self.total_targets}.")
                continue

        d.add_page_break()
        print("[+] Finished writing Detailed Analysis section...")
        # End of the detailed results and the beginning of browser, location, and OS stats
        add_heading("Statistiken")

        # Erstellung der Listen für Plot
        # Browser Familien
        counted_browser_family = Counter(self.browsers_family)
        browser_family_label = []
        browser_family_data = []
        for key, value in counted_browser_family.items():
            browser_family_label.append(key)
            browser_family_data.append(value)
        create_fig(browser_family_label,browser_family_data,'Browser')
        # OS Familien
        counted_operating_systems_family = Counter(self.operating_systems_family)
        operating_systems_label = []
        operating_systems_data = []
        for key, value in counted_operating_systems_family.items():
            operating_systems_label.append(key)
            operating_systems_data.append(value)
        create_fig(operating_systems_label,operating_systems_data,'Betriebssysteme')
        # hinzufügen der Diagramme
        table = d.add_table(1, 2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        temp_cell = table.cell(0, 0)
        temp_cell.text=""
        temp_cell.paragraphs[0].add_run().add_picture('./files/Browser.png')
        temp_cell = table.cell(0, 1)
        temp_cell.text=""
        temp_cell.paragraphs[0].add_run().add_picture('./files/Betriebssysteme.png')
        # Löschen der Bilder
        os.remove('./files/Browser.png')
        os.remove('./files/Betriebssysteme.png')

        # hinzufügen des Timeline_Chart
        time_chart()
        d.add_paragraph().add_run().add_picture('./files/timeline.png',width=Inches(6.8))
        os.remove('./files/timeline.png')

        d.add_page_break()

        add_heading("Detailierte Statistiken")

        p = d.add_paragraph("Die folgende Tabelle zeigt eine Übersicht der Browser:")
        # Create browser table
        browser_table = d.add_table(rows=1, cols=2, style=table_style)
        self._set_word_column_width(browser_table.columns[0], Cm(7.24))
        self._set_word_column_width(browser_table.columns[1], Cm(3.35))

        header1 = browser_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Browser", "Cell Text").bold = True

        header2 = browser_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Häufigkeit", "Cell Text").bold = True

        p = d.add_paragraph("\nDie folgende Tabelle zeigt die erfassten Betriebssysteme:")

        # Create OS table
        os_table = d.add_table(rows=1, cols=2, style=table_style)
        self._set_word_column_width(os_table.columns[0], Cm(7.24))
        self._set_word_column_width(os_table.columns[1], Cm(3.35))

        header1 = os_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Betriebssystem", "Cell Text").bold = True

        header2 = os_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Häufigkeit", "Cell Text").bold = True

        p = d.add_paragraph("\nDie folgende Tabelle zeigt die erfassten Positionen:")

        # Create geo IP table
        location_table = d.add_table(rows=1, cols=2, style=table_style)
        self._set_word_column_width(location_table.columns[0], Cm(7.24))
        self._set_word_column_width(location_table.columns[1], Cm(3.35))

        header1 = location_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Position", "Cell Text").bold = True

        header2 = location_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Besuche", "Cell Text").bold = True

        p = d.add_paragraph("\nDie folgende Tabelle zeigt die erfassten IP-Adressen:")

        # Create IP address table
        ip_add_table = d.add_table(rows=1, cols=2, style=table_style)
        self._set_word_column_width(ip_add_table.columns[0], Cm(7.24))
        self._set_word_column_width(ip_add_table.columns[1], Cm(3.35))

        header1 = ip_add_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("IP Adresse", "Cell Text").bold = True

        header2 = ip_add_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Häufigkeit", "Cell Text").bold = True

        p = d.add_paragraph("\nDie folgende Tabelle zeigt die mit Geolokalisierungsdaten abgeglichenen IP-Adressen:")

        # Create IP address and location table
        ip_loc_table = d.add_table(rows=1, cols=2, style=table_style)
        self._set_word_column_width(ip_loc_table.columns[0], Cm(7.24))
        self._set_word_column_width(ip_loc_table.columns[1], Cm(3.35))

        header1 = ip_loc_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("IP Adresse", "Cell Text").bold = True

        header2 = ip_loc_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Position", "Cell Text").bold = True

        # Counters are used here again to track rows
        counter = 1
        # Counter is used to count all elements in the lists to create a unique list with totals
        counted_browsers = Counter(self.browsers)
        counted_browsers = collections.OrderedDict(sorted(counted_browsers.items()))
        for key, value in counted_browsers.items():
            browser_table.add_row()
            cell = browser_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = browser_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        counted_os = Counter(self.operating_systems)
        counted_os = collections.OrderedDict(sorted(counted_os.items()))
        for key, value in counted_os.items():
            os_table.add_row()
            cell = os_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = os_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        counted_locations = Counter(self.locations)
        for key, value in counted_locations.items():
            location_table.add_row()
            cell = location_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = location_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        counted_ip_addresses = Counter(self.ip_addresses)
        for key, value in counted_ip_addresses.items():
            ip_add_table.add_row()
            cell = ip_add_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = ip_add_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        counter = 1
        for key, value in self.ip_and_location.items():
            ip_loc_table.add_row()
            cell = ip_loc_table.cell(counter, 0)
            cell.text = f"{key}"

            cell = ip_loc_table.cell(counter, 1)
            cell.text = f"{value}"
            counter += 1

        # Finalize document and save it as the value of output_word_report
        d.save(f"{self.output_word_report}")
        print(f"[+] Done! Check \"{self.output_word_report}\" for your results.")



    def config_section_map(self, config_parser, section):
        """This function helps by reading accepting a config file section, from gophish.config,
        and returning a dictionary object that can be referenced for configuration settings.
        """
        section_dict = {}
        options = config_parser.options(section)
        for option in options:
            try:
                section_dict[option] = config_parser.get(section, option)
                if section_dict[option] == -1:
                    print(f"[-] Skipping: {option}")
            except:
                print(f"[!] There was an error with: {option}")
                section_dict[option] = None
        return section_dict
